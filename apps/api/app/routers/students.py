"""Students CRUD — admin + super_admin can manage students."""

import csv
import io
import uuid

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from pydantic import BaseModel
from sqlalchemy import select, func, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.deps import get_db
from app.auth.deps import get_current_user
from app.auth.security import hash_password
from app.models.user import User, UserRole
from app.models.student import Student

router = APIRouter(prefix="/v1/students", tags=["students"])

DEFAULT_PASSWORD = "querion123"


# -- Auth dependency: admin or super_admin --
async def require_admin(user: User = Depends(get_current_user)) -> User:
    if user.role not in (UserRole.super_admin, UserRole.admin):
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


# -- Schemas --
class CreateStudentRequest(BaseModel):
    email: str
    name: str
    student_id: str | None = None
    password: str | None = None  # optional, defaults to DEFAULT_PASSWORD


class UpdateStudentRequest(BaseModel):
    name: str | None = None
    email: str | None = None
    student_id: str | None = None
    is_active: bool | None = None


class StudentItem(BaseModel):
    id: str
    email: str
    name: str
    student_id: str | None
    is_active: bool
    must_change_password: bool
    created_at: str


class ImportResult(BaseModel):
    created: int
    skipped: int
    errors: list[str]


def _to_item(s: Student) -> StudentItem:
    return StudentItem(
        id=str(s.id),
        email=s.email,
        name=s.name,
        student_id=s.student_id,
        is_active=s.is_active,
        must_change_password=s.must_change_password,
        created_at=s.created_at.isoformat(),
    )


# -- Endpoints --
@router.post("", response_model=StudentItem, status_code=201)
async def create_student(
    body: CreateStudentRequest,
    _: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    existing = await db.execute(select(Student).where(Student.email == body.email))
    if existing.scalar_one_or_none():
        raise HTTPException(status_code=409, detail="Email already registered")

    pwd = body.password or DEFAULT_PASSWORD
    student = Student(
        email=body.email,
        password_hash=hash_password(pwd),
        name=body.name,
        student_id=body.student_id,
        must_change_password=(pwd == DEFAULT_PASSWORD),
    )
    db.add(student)
    await db.commit()
    await db.refresh(student)
    return _to_item(student)


@router.get("", response_model=list[StudentItem])
async def list_students(
    _: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(Student).order_by(Student.created_at.desc()))
    return [_to_item(s) for s in result.scalars().all()]


@router.post("/import", response_model=ImportResult)
async def import_students(
    file: UploadFile = File(...),
    _: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """Import students from CSV, Excel (.xlsx), or Word (.docx).
    Word files should contain a table with headers: email, name, student_id.
    """
    filename = file.filename or ""
    content = await file.read()
    rows_data: list[dict[str, str]] = []

    def map_headers(headers: list[str]) -> dict[int, str]:
        mapping = {}
        email_keywords = ["email", "thư điện tử"]
        name_keywords = ["name", "tên", "họ tên", "họ và tên"]
        sid_keywords = ["student_id", "mssv", "mã sinh viên", "mã số sinh viên", "id"]

        for i, h in enumerate(headers):
            h_lower = h.lower()
            if any(k in h_lower for k in email_keywords):
                mapping[i] = "email"
            elif any(k in h_lower for k in name_keywords):
                mapping[i] = "name"
            elif any(k in h_lower for k in sid_keywords):
                mapping[i] = "student_id"
        return mapping

    try:
        if filename.endswith(".csv"):
            text = content.decode("utf-8-sig")
            reader = csv.reader(io.StringIO(text))
            rows = [r for r in reader if any(cell.strip() for cell in r)]
            if not rows:
                raise ValueError("CSV file is empty")
            
            # Find header row
            header_idx = 0
            for idx, r in enumerate(rows):
                if any(k in r[0].lower() or (len(r)>1 and k in r[1].lower()) for k in ["email", "name", "tên"]):
                    header_idx = idx
                    break
            
            headers = [h.strip() for h in rows[header_idx]]
            mapping = map_headers(headers)
            
            for r in rows[header_idx + 1:]:
                row_dict = {}
                for idx, col_name in mapping.items():
                    if idx < len(r):
                        row_dict[col_name] = r[idx].strip()
                if row_dict:
                    rows_data.append(row_dict)

        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(content), data_only=True)
            sheet = wb.active
            if not sheet:
                raise ValueError("Excel file has no active sheet")
            
            all_rows = list(sheet.iter_rows(values_only=True))
            # Skip empty rows at start
            start_row = 0
            for i, r in enumerate(all_rows):
                if any(cell is not None and str(cell).strip() for cell in r):
                    start_row = i
                    break
            
            headers = [str(h).strip() if h is not None else "" for h in all_rows[start_row]]
            mapping = map_headers(headers)
            
            for r in all_rows[start_row + 1:]:
                row_dict = {}
                for idx, col_name in mapping.items():
                    if idx < len(r):
                        val = r[idx]
                        row_dict[col_name] = str(val).strip() if val is not None else ""
                if any(row_dict.values()):
                    rows_data.append(row_dict)

        elif filename.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            if not doc.tables:
                raise ValueError("Word document contains no tables")
            
            table = doc.tables[0]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            mapping = map_headers(headers)
            
            for row in table.rows[1:]:
                row_dict = {}
                for idx, col_name in mapping.items():
                    if idx < len(row.cells):
                        row_dict[col_name] = row.cells[idx].text.strip()
                if any(row_dict.values()):
                    rows_data.append(row_dict)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use CSV, Excel, or Word.")

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

    created = 0
    skipped = 0
    errors: list[str] = []
    hashed_default = hash_password(DEFAULT_PASSWORD)

    for i, row in enumerate(rows_data, start=2):
        email = (row.get("email") or "").strip()
        name = (row.get("name") or "").strip()
        sid = (row.get("student_id") or "").strip() or None

        if not email or not name:
            errors.append(f"Row {i}: missing email or name")
            continue

        existing = await db.execute(select(Student.id).where(Student.email == email))
        if existing.scalar_one_or_none():
            skipped += 1
            continue

        db.add(Student(
            email=email,
            password_hash=hashed_default,
            name=name,
            student_id=sid,
            must_change_password=True,
        ))
        created += 1

    await db.commit()
    return ImportResult(created=created, skipped=skipped, errors=errors)


@router.patch("/{student_id}", response_model=StudentItem)
async def update_student(
    student_id: str,
    body: UpdateStudentRequest,
    _: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    student = await db.get(Student, uuid.UUID(student_id))
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    if body.name is not None:
        student.name = body.name
    if body.email is not None:
        # Check email uniqueness if changing
        if body.email != student.email:
            existing = await db.execute(select(Student).where(Student.email == body.email))
            if existing.scalar_one_or_none():
                raise HTTPException(status_code=409, detail="Email already registered")
        student.email = body.email
    if body.student_id is not None:
        student.student_id = body.student_id
    if body.is_active is not None:
        student.is_active = body.is_active

    await db.commit()
    await db.refresh(student)
    return _to_item(student)


@router.delete("/{student_id}", status_code=200)
async def delete_student(
    student_id: str,
    _: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """Permanently delete a student."""
    sid_uuid = uuid.UUID(student_id)
    student = await db.get(Student, sid_uuid)
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    
    await db.execute(delete(Student).where(Student.id == sid_uuid))
    await db.commit()
    return {"detail": "Student deleted"}
