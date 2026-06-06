import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_specification():
    doc = Document()
    
    # Title
    title = doc.add_heading('ĐẶC TẢ YÊU CẦU DỰ ÁN QUERION (MINI-DIFY)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Tổng quan
    doc.add_heading('1. Tổng quan dự án', level=1)
    doc.add_paragraph(
        'Querion là một nền tảng quản lý tri thức và xây dựng ứng dụng AI (Mini-Dify). '
        'Hệ thống cho phép người dùng xây dựng các kho tri thức từ tài liệu cá nhân, '
        'thiết kế các quy trình xử lý (workflow) linh hoạt và triển khai các chatbot AI '
        'có khả năng trích dẫn nguồn dữ liệu chính xác.'
    )
    
    # 2. Đối tượng người dùng
    doc.add_heading('2. Đối tượng người dùng và Vai trò', level=1)
    
    doc.add_heading('2.1. Super Admin (Quản trị tối cao)', level=2)
    doc.add_paragraph('Có quyền kiểm soát toàn bộ hệ thống:', style='List Bullet')
    doc.add_paragraph('Quản lý danh sách Workspace.', style='List Bullet')
    doc.add_paragraph('Quản lý người dùng cấp Admin.', style='List Bullet')
    doc.add_paragraph('Cấu hình các nhà cung cấp AI (OpenAI, Anthropic, Google, v.v.).', style='List Bullet')
    
    doc.add_heading('2.2. Admin (Quản trị Workspace)', level=2)
    doc.add_paragraph('Được phân quyền vào các Workspace cụ thể:', style='List Bullet')
    doc.add_paragraph('Owner: Toàn quyền trong Workspace, quản lý thành viên.', style='List Bullet')
    doc.add_paragraph('Editor: Quản lý Dataset, Workflow, App nhưng không quản lý được thành viên.', style='List Bullet')
    doc.add_paragraph('Viewer: Chỉ xem và sử dụng các tài nguyên trong Workspace.', style='List Bullet')
    
    doc.add_heading('2.3. Student (Người dùng cuối)', level=2)
    doc.add_paragraph('Truy cập vào các ứng dụng đã được phát bản (publish).', style='List Bullet')
    doc.add_paragraph('Thực hiện chat với AI dựa trên tri thức đã cấu hình.', style='List Bullet')
    doc.add_paragraph('Quản lý lịch sử trò chuyện cá nhân.', style='List Bullet')
    
    # 3. Yêu cầu chức năng
    doc.add_heading('3. Yêu cầu chức năng chính', level=1)
    
    doc.add_heading('3.1. Quản lý tri thức (Dataset & Knowledge Base)', level=2)
    doc.add_paragraph('Tải lên tài liệu (PDF, Word, TXT, Excel, CSV).', style='List Bullet')
    doc.add_paragraph('Tự động phân tách văn bản (Chunking) và đánh chỉ mục vector.', style='List Bullet')
    doc.add_paragraph('Quản lý trạng thái xử lý tài liệu.', style='List Bullet')
    
    doc.add_heading('3.2. Thiết kế luồng (Workflow Canvas)', level=2)
    doc.add_paragraph('Giao diện kéo thả (Flow-based) để thiết kế quy trình AI.', style='List Bullet')
    doc.add_paragraph('Hỗ trợ nhiều loại node: Input, Retrieval, Prompt, LLM, Condition, Output.', style='List Bullet')
    doc.add_paragraph('Lưu trữ và kiểm tra tính hợp lệ của workflow.', style='List Bullet')
    
    doc.add_heading('3.3. Ứng dụng & Phát bản (App Management)', level=2)
    doc.add_paragraph('Tạo ứng dụng chat dựa trên Workflow hoặc Dataset.', style='List Bullet')
    doc.add_paragraph('Phát bản ứng dụng để sinh viên có thể truy cập.', style='List Bullet')
    doc.add_paragraph('Cấu hình các tham số AI và System Prompt.', style='List Bullet')
    
    doc.add_heading('3.4. Trò chuyện & Trích dẫn (Chat & Citation)', level=2)
    doc.add_paragraph('Chat streaming (phản hồi theo thời gian thực).', style='List Bullet')
    doc.add_paragraph('Trích dẫn nguồn tài liệu cụ thể cho từng câu trả lời.', style='List Bullet')
    doc.add_paragraph('Hỗ trợ hội thoại đa bước (Context-aware).', style='List Bullet')
    
    # 4. Kiến trúc & Công nghệ
    doc.add_heading('4. Kiến trúc hệ thống và Công nghệ', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Thành phần'
    hdr_cells[1].text = 'Công nghệ'
    
    tech_stack = [
        ('Frontend', 'Next.js 15, TypeScript, Tailwind CSS v4, XYFlow'),
        ('Backend API', 'FastAPI (Python 3.11+), LangGraph'),
        ('Database', 'PostgreSQL + pgvector (Vector Store)'),
        ('Cache & Queue', 'Redis'),
        ('Object Storage', 'MinIO (S3 Compatible)'),
        ('Worker', 'Background Indexing Service')
    ]
    
    for comp, tech in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    # 5. Yêu cầu phi chức năng
    doc.add_heading('5. Yêu cầu phi chức năng', level=1)
    doc.add_paragraph('Bảo mật: Mã hóa API Key, phân quyền chặt chẽ theo Workspace.', style='List Bullet')
    doc.add_paragraph('Hiệu năng: Phản hồi streaming dưới 500ms, tìm kiếm vector tối ưu.', style='List Bullet')
    doc.add_paragraph('Khả năng mở rộng: Thiết kế dạng Microservices, sẵn sàng cho Docker/K8s.', style='List Bullet')

    # Save
    filename = 'Querion_Specification.docx'
    doc.save(filename)
    print(f"File saved: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_specification()
