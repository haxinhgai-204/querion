import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_detailed_spec():
    doc = Document()
    
    # Title
    title = doc.add_heading('ĐẶC TẢ CHỨC NĂNG CHI TIẾT HỆ THỐNG QUERION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Quản lý Workspace
    doc.add_heading('1. Phân hệ Quản lý Workspace', level=1)
    doc.add_paragraph('Mục tiêu: Phân tách dữ liệu và tài nguyên giữa các nhóm người dùng khác nhau.')
    
    doc.add_heading('1.1. Quản lý danh sách Workspace (Super Admin)', level=2)
    doc.add_paragraph('Tạo mới Workspace: Nhập tên và mô tả.', style='List Bullet')
    doc.add_paragraph('Chỉnh sửa Workspace: Cập nhật thông tin cơ bản.', style='List Bullet')
    doc.add_paragraph('Xóa Workspace: Xóa toàn bộ dữ liệu liên quan (Dataset, App, Workflow) thông qua cơ chế Cascade Delete.', style='List Bullet')
    
    doc.add_heading('1.2. Quản lý thành viên Workspace (Owner)', level=2)
    doc.add_paragraph('Mời thành viên: Thêm các Admin khác vào Workspace.', style='List Bullet')
    doc.add_paragraph('Phân vai trò (Roles):', style='List Bullet')
    doc.add_paragraph('- Owner: Toàn quyền, bao gồm quản lý thành viên.', style='List Bullet')
    doc.add_paragraph('- Editor: Có quyền tạo/sửa/xóa Dataset, Workflow, App.', style='List Bullet')
    doc.add_paragraph('- Viewer: Chỉ có quyền xem và chạy thử ứng dụng.', style='List Bullet')
    
    # 2. Quản lý Tri thức (Dataset)
    doc.add_heading('2. Phân hệ Quản lý Tri thức (Dataset)', level=1)
    doc.add_paragraph('Mục tiêu: Xây dựng kho tri thức riêng từ các tài liệu nghiệp vụ.')
    
    doc.add_heading('2.1. Quản lý Dataset', level=2)
    doc.add_paragraph('Tạo Dataset: Đặt tên và chọn mục đích sử dụng.', style='List Bullet')
    doc.add_paragraph('Xóa Dataset: Gỡ bỏ toàn bộ tài liệu và vector liên quan.', style='List Bullet')
    
    doc.add_heading('2.2. Quản lý Tài liệu (Document)', level=2)
    doc.add_paragraph('Tải lên đa định dạng: Hỗ trợ PDF, Word, TXT, Excel, CSV.', style='List Bullet')
    doc.add_paragraph('Quy trình xử lý (Indexing Pipeline):', style='List Bullet')
    doc.add_paragraph('- Parsing: Trích xuất văn bản từ file.', style='List Bullet')
    doc.add_paragraph('- Chunking: Chia nhỏ văn bản theo kích thước cấu hình (ví dụ: 1000 tokens).', style='List Bullet')
    doc.add_paragraph('- Embedding: Chuyển đổi văn bản thành vector bằng các model (OpenAI, Google).', style='List Bullet')
    doc.add_paragraph('- Vector Store: Lưu trữ vào Postgres (pgvector) để tìm kiếm ngữ nghĩa.', style='List Bullet')
    doc.add_paragraph('Trạng thái tài liệu: Theo dõi quá trình xử lý (Uploaded, Indexing, Ready, Error).', style='List Bullet')
    
    # 3. Thiết kế luồng xử lý (Workflow Canvas)
    doc.add_heading('3. Phân hệ Thiết kế Luồng (Workflow)', level=1)
    doc.add_paragraph('Mục tiêu: Tùy biến quy trình xử lý của AI thông qua giao diện trực quan.')
    
    doc.add_heading('3.1. Giao diện Canvas', level=2)
    doc.add_paragraph('Sử dụng cơ chế Kéo-Thả (Drag & Drop) để kết nối các Node.', style='List Bullet')
    doc.add_paragraph('Tự động kiểm tra tính hợp lệ của luồng (Validation): Phải có 1 Node đầu vào và ít nhất 1 Node đầu ra.', style='List Bullet')
    
    doc.add_heading('3.2. Danh sách các Node chức năng', level=2)
    doc.add_paragraph('Input Node: Điểm bắt đầu, nhận câu hỏi từ người dùng.', style='List Bullet')
    doc.add_paragraph('Knowledge Retrieval Node: Tìm kiếm thông tin liên quan từ Dataset đã chọn.', style='List Bullet')
    doc.add_paragraph('LLM Node: Cấu hình Model, System Prompt và đưa Context vào Prompt.', style='List Bullet')
    doc.add_paragraph('Template Node: Soạn thảo văn bản tĩnh kết hợp biến động từ các node trước.', style='List Bullet')
    doc.add_paragraph('If-Else Node: Rẽ nhánh xử lý dựa trên điều kiện logic.', style='List Bullet')
    doc.add_paragraph('Output Node: Điểm kết thúc, trả về câu trả lời cho người dùng.', style='List Bullet')
    
    # 4. Quản lý Ứng dụng (App)
    doc.add_heading('4. Phân hệ Quản lý Ứng dụng', level=1)
    doc.add_paragraph('Mục tiêu: Đóng gói tri thức và quy trình thành ứng dụng chatbot hoàn chỉnh.')
    
    doc.add_heading('4.1. Cấu hình App', level=2)
    doc.add_paragraph('Chế độ Chat: Chat trực tiếp dựa trên 1 Dataset.', style='List Bullet')
    doc.add_paragraph('Chế độ Workflow: Chat dựa trên quy trình đã thiết kế ở Canvas.', style='List Bullet')
    doc.add_paragraph('Cấu hình tham số: Model LLM, Temperature, Max Tokens.', style='List Bullet')
    
    doc.add_heading('4.2. Phát bản (Publishing)', level=2)
    doc.add_paragraph('Tính năng Publish: Cho phép sinh viên nhìn thấy và sử dụng App.', style='List Bullet')
    doc.add_paragraph('Tạo API Key: Cung cấp khóa truy cập cho tích hợp bên thứ ba.', style='List Bullet')
    
    # 5. Giao diện Người dùng cuối (Student Interface)
    doc.add_heading('5. Giao diện dành cho Sinh viên', level=1)
    
    doc.add_heading('5.1. Chat & Trải nghiệm AI', level=2)
    doc.add_paragraph('Streaming Response: Câu trả lời hiển thị dần theo thời gian thực (Server-Sent Events).', style='List Bullet')
    doc.add_paragraph('Citations (Trích dẫn): Hiển thị rõ nguồn dữ liệu (tên file, nội dung đoạn trích) AI đã dùng để trả lời.', style='List Bullet')
    doc.add_paragraph('Lịch sử hội thoại: Tự động lưu và cho phép xem lại các cuộc trò chuyện cũ.', style='List Bullet')
    
    doc.add_heading('5.2. Quản lý tài khoản', level=2)
    doc.add_paragraph('Đăng nhập bằng Email/MSSV.', style='List Bullet')
    doc.add_paragraph('Yêu cầu đổi mật khẩu trong lần đăng nhập đầu tiên.', style='List Bullet')
    
    # 6. Quản trị hệ thống (Super Admin)
    doc.add_heading('6. Phân hệ Quản trị hệ thống', level=1)
    
    doc.add_heading('6.1. Quản lý Sinh viên', level=2)
    doc.add_paragraph('Import Sinh viên: Hỗ trợ nạp hàng loạt từ file CSV, Excel, Word.', style='List Bullet')
    doc.add_paragraph('Xử lý tiêu đề thông minh: Tự động nhận diện cột (Họ tên, Email, MSSV).', style='List Bullet')
    doc.add_paragraph('Kích hoạt/Vô hiệu hóa: Quản lý trạng thái hoạt động của sinh viên.', style='List Bullet')
    
    doc.add_heading('6.2. Quản lý AI Providers', level=2)
    doc.add_paragraph('Cấu hình API Key: OpenAI, Anthropic, Google Gemini, OpenRouter.', style='List Bullet')
    doc.add_paragraph('Bảo mật: API Key được mã hóa AES-256 trước khi lưu vào cơ sở dữ liệu.', style='List Bullet')

    # Save
    filename = 'Querion_Detailed_Functional_Spec.docx'
    doc.save(filename)
    print(f"File saved: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_detailed_spec()
